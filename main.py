import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE

# Open the PDF
pdf_path = "input_files/input.pdf"
docx_path = "output_files/output.docx"
pdf_document = fitz.open(pdf_path)

# Create a new DOCX document
doc = Document()


# Function to add hyperlinks to DOCX
def add_hyperlink(paragraph, url, text):
    # Create a hyperlink in the paragraph
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    new_run.append(rPr)
    new_run.text = text

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# Iterate through each page
for page_num in range(len(pdf_document)):
    page = pdf_document.load_page(page_num)

    # Extract plain text from the PDF
    text = page.get_text("text")
    doc.add_paragraph(text)

    # Handle images
    images = page.get_images(full=True)
    for img_index, img in enumerate(images):
        xref = img[0]
        base_image = pdf_document.extract_image(xref)
        image_bytes = base_image["image"]

        # Save the image and add it to the DOCX
        image_filename = f"image_{page_num}_{img_index}.png"
        with open(image_filename, "wb") as img_file:
            img_file.write(image_bytes)

        doc.add_picture(image_filename, width=Inches(6))  # Add image to DOCX

    # Handle hyperlinks and annotations
    links = page.get_links()
    for link in links:
        if "uri" in link:
            uri = link["uri"]
            rect = link["from"]  # The rectangle area of the link
            link_text = page.get_text("text", clip=rect)
            if link_text:
                paragraph = doc.add_paragraph()  # Create a new paragraph for the link
                add_hyperlink(paragraph, uri, link_text)

# Save the DOCX file
doc.save(docx_path)
print(f"Conversion complete: {docx_path}")
